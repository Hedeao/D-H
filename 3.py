import os
import random
from docx import Document

def is_prime(num):
    if num <= 1:
        return False
    for i in range(2, int(num**0.5) + 1):
        if num % i == 0:
            return False
    return True

def find_primitive_root(p):
    required_set = set(num for num in range(1, p))
    for g in range(2, p):
        actual_set = set(pow(g, powers, p) for powers in range(1, p))
        if required_set == actual_set:
            return g
    return None

# Step 1: Choose a prime number p and a primitive root g
while True:
    p = random.randint(100, 255)
    if is_prime(p):
        break

g = find_primitive_root(p)

# Create a Word document
doc = Document()
doc.add_heading('Diffie-Hellman 密钥交换机制验证', level=1)

doc.add_paragraph(f"选择的素数 (p): {p}")
doc.add_paragraph(f"选择的原根 (g): {g}")

# Step 2: A and B choose private keys
a_private_key = random.randint(2, p - 2)
b_private_key = random.randint(2, p - 2)

doc.add_paragraph(f"A 的私钥 (a): {a_private_key}")
doc.add_paragraph(f"B 的私钥 (b): {b_private_key}")

# Step 3: A and B compute public keys
a_public_key = pow(g, a_private_key, p)
b_public_key = pow(g, b_private_key, p)

doc.add_paragraph(f"A 计算公钥 (A = g^a mod p):")
doc.add_paragraph(f"A = {g}^{a_private_key} mod {p} = {a_public_key}")
doc.add_paragraph(f"B 计算公钥 (B = g^b mod p):")
doc.add_paragraph(f"B = {g}^{b_private_key} mod {p} = {b_public_key}")

# Step 4: A and B compute shared secret key
a_shared_secret = pow(b_public_key, a_private_key, p)
b_shared_secret = pow(a_public_key, b_private_key, p)

doc.add_paragraph(f"A 计算共享秘密密钥 (s_A = B^a mod p):")
doc.add_paragraph(f"s_A = {b_public_key}^{a_private_key} mod {p} = {a_shared_secret}")
doc.add_paragraph(f"B 计算共享秘密密钥 (s_B = A^b mod p):")
doc.add_paragraph(f"s_B = {a_public_key}^{b_private_key} mod {p} = {b_shared_secret}")

# Verify that both parties have the same shared secret key
assert a_shared_secret == b_shared_secret, "共享秘密密钥不匹配!"

doc.add_paragraph("验证成功：A 和 B 得到了相同的共享秘密密钥。")

# Define the save path
save_path = r'E:\20221401赵浚杰\xinandaolun\dh_key_exchange.docx'

# Ensure the directory exists
os.makedirs(os.path.dirname(save_path), exist_ok=True)

# Save the document
doc.save(save_path)








